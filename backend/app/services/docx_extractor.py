import docx
from typing import Dict, Any, List
from loguru import logger
import io

class DOCXExtractor:
    def __init__(self):
        pass

    async def extract(self, file_content: bytes) -> Dict[str, Any]:
        """
        Extracts text, metadata, and links from DOCX content.
        """
        raw_text = ""
        metadata = {
            "file_type": "docx",
            "properties": {}
        }
        links = []

        try:
            doc = docx.Document(io.BytesIO(file_content))
            
            # 1. Extract text from paragraphs
            for para in doc.paragraphs:
                raw_text += para.text + "\n"
                
                # Extract hyperlinks (docx-python requires looking at XML for this usually)
                # But we can also rely on our LinkExtractor (regex) for DOCX 
                # as docx-python's hyperlink support is limited.
            
            # 2. Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        raw_text += cell.text + " "
                    raw_text += "\n"

            # 3. Extract metadata
            props = doc.core_properties
            metadata["properties"] = {
                "author": props.author,
                "created": str(props.created) if props.created else None,
                "modified": str(props.modified) if props.modified else None,
                "title": props.title
            }
            
            logger.info("Successfully extracted DOCX")
            
            return {
                "raw_text": raw_text,
                "metadata": metadata,
                "links": [] # Will rely on LinkExtractor regex for DOCX
            }

        except Exception as e:
            logger.error(f"Error extracting DOCX: {e}")
            raise Exception(f"DOCX extraction failed: {str(e)}")
